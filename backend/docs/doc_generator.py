"""
Documentation Generator with Multiple Output Formats and Styling

Comprehensive documentation generation for RaptorFlow backend:
- Multiple output formats (HTML, PDF, Markdown, Word)
- Customizable styling and themes
- Interactive documentation features
- Code examples and tutorials
- API reference generation
"""

import asyncio
import json
import logging
import os
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum

import jinja2
import markdown
import pdfkit
import yaml
from docx import Document
from jinja2 import Environment, FileSystemLoader
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)


class OutputFormat(Enum):
    """Documentation output formats."""
    HTML = "html"
    PDF = "pdf"
    MARKDOWN = "md"
    WORD = "docx"
    JSON = "json"


class DocumentationTheme(Enum):
    """Documentation themes."""
    DEFAULT = "default"
    DARK = "dark"
    MODERN = "modern"
    TECHNICAL = "technical"
    MINIMAL = "minimal"


@dataclass
class DocumentationSection:
    """Documentation section configuration."""
    title: str
    content: str
    subsections: List['DocumentationSection'] = field(default_factory=list)
    code_examples: List[Dict[str, str]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentationConfig(BaseModel):
    """Documentation generation configuration."""
    openapi_spec_path: str = "docs/openapi_comprehensive.yaml"
    output_dir: str = "generated_docs"
    formats: List[OutputFormat] = Field(default=[OutputFormat.HTML])
    theme: DocumentationTheme = DocumentationTheme.DEFAULT
    include_code_examples: bool = True
    include_tutorials: bool = True
    include_api_reference: bool = True
    custom_css: Optional[str] = None
    custom_js: Optional[str] = None
    logo_url: Optional[str] = None
    company_name: str = "RaptorFlow"
    product_name: str = "Backend API"
    version: str = "1.0.0"


class DocumentationGenerator:
    """Comprehensive documentation generator."""

    def __init__(self, config: DocumentationConfig):
        self.config = config
        self.openapi_spec: Dict[str, Any] = {}
        self.template_env: Optional[jinja2.Environment] = None
        
        # Create output directory
        Path(config.output_dir).mkdir(parents=True, exist_ok=True)
        
        # Load OpenAPI specification
        self._load_openapi_spec()
        
        # Setup Jinja2 environment
        self._setup_template_env()

    def _load_openapi_spec(self) -> None:
        """Load OpenAPI specification."""
        try:
            with open(self.config.openapi_spec_path, 'r') as f:
                self.openapi_spec = yaml.safe_load(f)
            logger.info(f"Loaded OpenAPI spec: {self.config.openapi_spec_path}")
        except Exception as e:
            logger.error(f"Failed to load OpenAPI spec: {e}")
            raise

    def _setup_template_env(self) -> None:
        """Setup Jinja2 template environment."""
        template_dir = Path(__file__).parent / "templates"
        self.template_env = Environment(
            loader=FileSystemLoader([template_dir]),
            autoescape=jinja2.select_autoescape(['html', 'xml'])
        )
        
        # Add custom filters
        self.template_env.filters['markdown'] = self._markdown_filter
        self.template_env.filters['code_highlight'] = self._code_highlight_filter
        self.template_env.filters['format_date'] = self._format_date_filter

    def _markdown_filter(self, text: str) -> str:
        """Convert markdown to HTML."""
        return markdown.markdown(text, extensions=['codehilite', 'tables', 'toc'])

    def _code_highlight_filter(self, code: str, language: str = 'python') -> str:
        """Apply syntax highlighting to code."""
        try:
            import pygments
            from pygments import highlight
            from pygments.lexers import get_lexer_by_name
            from pygments.formatters import HtmlFormatter
            
            lexer = get_lexer_by_name(language)
            formatter = HtmlFormatter(style='github', cssclass='highlight')
            return highlight(code, lexer, formatter)
        except ImportError:
            return f'<pre><code class="language-{language}">{code}</code></pre>'

    def _format_date_filter(self, date_str: str) -> str:
        """Format date string."""
        try:
            date_obj = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
            return date_obj.strftime('%B %d, %Y')
        except:
            return date_str

    def generate_tutorial_sections(self) -> List[DocumentationSection]:
        """Generate tutorial sections."""
        tutorials = [
            DocumentationSection(
                title="Getting Started",
                content="""
                Welcome to the RaptorFlow Backend API documentation. This guide will help you get started
                with integrating our powerful AI-powered business intelligence platform into your applications.
                """,
                subsections=[
                    DocumentationSection(
                        title="Authentication",
                        content="""
                        All API endpoints require authentication using JWT tokens. First, obtain an access token
                        by logging in with your credentials.
                        """,
                        code_examples=[
                            {
                                "language": "python",
                                "code": '''
import requests

# Login to get access token
response = requests.post('https://api.raptorflow.com/auth/login', json={
    'email': 'your-email@example.com',
    'password': 'your-password'
})

if response.status_code == 200:
    token = response.json()['data']['access_token']
    headers = {'Authorization': f'Bearer {token}'}
    # Use token for authenticated requests
                                '''
                            },
                            {
                                "language": "javascript",
                                "code": '''
// Login to get access token
const response = await fetch('https://api.raptorflow.com/auth/login', {
  method: 'POST',
  headers: {'Content-Type': 'application/json'},
  body: JSON.stringify({
    email: 'your-email@example.com',
    password: 'your-password'
  })
});

if (response.ok) {
  const data = await response.json();
  const token = data.data.access_token;
  const headers = {'Authorization': `Bearer ${token}`};
  // Use token for authenticated requests
}
                                '''
                            }
                        ]
                    ),
                    DocumentationSection(
                        title="Making Your First Request",
                        content="""
                        Once you have your access token, you can start making API requests.
                        Here's how to get your user profile information.
                        """,
                        code_examples=[
                            {
                                "language": "python",
                                "code": '''
import requests

# Get user profile
headers = {'Authorization': 'Bearer YOUR_ACCESS_TOKEN'}
response = requests.get('https://api.raptorflow.com/users/me', headers=headers)

if response.status_code == 200:
    user_data = response.json()['data']
    print(f"Welcome, {user_data['name']}!")
                                '''
                            }
                        ]
                    )
                ]
            ),
            DocumentationSection(
                title="Working with Workspaces",
                content="""
                Workspaces are the primary container for your business data and AI agents.
                Learn how to create, manage, and interact with workspaces.
                """,
                subsections=[
                    DocumentationSection(
                        title="Creating a Workspace",
                        content="""
                        Create a new workspace to organize your business intelligence activities.
                        """,
                        code_examples=[
                            {
                                "language": "python",
                                "code": '''
import requests

# Create a new workspace
headers = {'Authorization': 'Bearer YOUR_ACCESS_TOKEN'}
workspace_data = {
    'name': 'My Business Workspace',
    'description': 'Workspace for business intelligence',
    'subscription_tier': 'pro'
}

response = requests.post(
    'https://api.raptorflow.com/workspaces',
    json=workspace_data,
    headers=headers
)

if response.status_code == 201:
    workspace = response.json()['data']
    print(f"Created workspace: {workspace['id']}")
                                '''
                            }
                        ]
                    ),
                    DocumentationSection(
                        title="Listing Workspaces",
                        content="""
                        Retrieve all workspaces associated with your account.
                        """,
                        code_examples=[
                            {
                                "language": "python",
                                "code": '''
import requests

# List all workspaces
headers = {'Authorization': 'Bearer YOUR_ACCESS_TOKEN'}
response = requests.get('https://api.raptorflow.com/workspaces', headers=headers)

if response.status_code == 200:
    workspaces = response.json()['data']['workspaces']
    for workspace in workspaces:
        print(f"Workspace: {workspace['name']} ({workspace['id']})")
                                '''
                            }
                        ]
                    )
                ]
            ),
            DocumentationSection(
                title="AI Agent Integration",
                content="""
                RaptorFlow's power comes from its AI agents. Learn how to execute agents
                for market research, content creation, and business analysis.
                """,
                subsections=[
                    DocumentationSection(
                        title="Executing Market Research Agent",
                        content="""
                        Use the market research agent to analyze market trends and competitive intelligence.
                        """,
                        code_examples=[
                            {
                                "language": "python",
                                "code": '''
import requests

# Execute market research agent
headers = {'Authorization': 'Bearer YOUR_ACCESS_TOKEN'}
agent_request = {
    'workspace_id': 'YOUR_WORKSPACE_ID',
    'input': 'Analyze market trends for SaaS companies in 2024',
    'context': {
        'industry': 'technology',
        'region': 'north_america',
        'timeframe': 'last_6_months'
    },
    'options': {
        'max_tokens': 1000,
        'temperature': 0.7,
        'include_sources': True
    }
}

response = requests.post(
    'https://api.raptorflow.com/agents/market_research/execute',
    json=agent_request,
    headers=headers
)

if response.status_code == 200:
    result = response.json()['data']
    print(f"Analysis: {result['output']}")
    print(f"Confidence: {result['confidence']}")
    print(f"Sources: {len(result['sources'])} found")
                                '''
                            }
                        ]
                    ),
                    DocumentationSection(
                        title="VertexAI Model Configuration",
                        content="""
                        RaptorFlow uses Google VertexAI as the primary AI provider. Configure
                        model parameters for optimal results.
                        """,
                        code_examples=[
                            {
                                "language": "python",
                                "code": '''
# VertexAI model configuration
model_config = {
    'model': 'gemini-pro',  # Primary VertexAI model
    'max_tokens': 2048,     # Maximum response length
    'temperature': 0.7,       # Creativity level (0.0-1.0)
    'top_p': 0.9,           # Nucleus sampling
    'top_k': 40              # Top-k sampling
}

# Use with AI inference endpoint
inference_request = {
    'prompt': 'Your business question here',
    **model_config
}
                                '''
                            }
                        ]
                    )
                ]
            )
        ]
        
        return tutorials

    def generate_api_reference_sections(self) -> List[DocumentationSection]:
        """Generate API reference sections."""
        if 'paths' not in self.openapi_spec:
            return []
        
        sections = []
        
        # Group endpoints by tag
        tagged_endpoints = {}
        for path, methods in self.openapi_spec['paths'].items():
            for method, details in methods.items():
                tags = details.get('tags', ['General'])
                for tag in tags:
                    if tag not in tagged_endpoints:
                        tagged_endpoints[tag] = []
                    tagged_endpoints[tag].append({
                        'path': path,
                        'method': method.upper(),
                        'details': details
                    })
        
        # Create sections for each tag
        for tag, endpoints in tagged_endpoints.items():
            subsections = []
            
            for endpoint in endpoints:
                path = endpoint['path']
                method = endpoint['method']
                details = endpoint['details']
                
                # Generate parameter documentation
                params_html = ""
                if 'parameters' in details:
                    params_html = "<table class='params-table'>"
                    params_html += "<thead><tr><th>Name</th><th>Type</th><th>Required</th><th>Description</th></tr></thead>"
                    params_html += "<tbody>"
                    
                    for param in details['parameters']:
                        params_html += f"""
                        <tr>
                            <td><code>{param['name']}</code></td>
                            <td>{param.get('schema', {}).get('type', 'string')}</td>
                            <td>{'Yes' if param.get('required', False) else 'No'}</td>
                            <td>{param.get('description', '')}</td>
                        </tr>
                        """
                    
                    params_html += "</tbody></table>"
                
                # Generate request body documentation
                body_html = ""
                if 'requestBody' in details:
                    body_schema = details['requestBody'].get('content', {}).get('application/json', {}).get('schema', {})
                    body_html = f"<pre><code>{json.dumps(body_schema, indent=2)}</code></pre>"
                
                # Generate response documentation
                responses_html = ""
                if 'responses' in details:
                    responses_html = "<div class='responses'>"
                    for status, response in details['responses'].items():
                        responses_html += f"""
                        <div class='response-item'>
                            <h4>Status {status}</h4>
                            <p>{response.get('description', '')}</p>
                            {f"<pre><code>{json.dumps(response.get('content', {}).get('application/json', {}).get('schema', {}), indent=2)}</code></pre>" if 'content' in response else ''}
                        </div>
                        """
                    responses_html += "</div>"
                
                # Create subsection
                subsection = DocumentationSection(
                    title=f"{method} {path}",
                    content=f"""
                    <div class="endpoint-description">
                        <p><strong>{details.get('summary', '')}</strong></p>
                        <p>{details.get('description', '')}</p>
                    </div>
                    
                    {f"<h3>Parameters</h3>{params_html}" if params_html else ''}
                    
                    {f"<h3>Request Body</h3>{body_html}" if body_html else ''}
                    
                    {f"<h3>Responses</h3>{responses_html}" if responses_html else ''}
                    """,
                    metadata={
                        'method': method,
                        'path': path,
                        'operation_id': details.get('operationId', '')
                    }
                )
                subsections.append(subsection)
            
            section = DocumentationSection(
                title=tag,
                content=f"API endpoints for {tag} operations.",
                subsections=subsections
            )
            sections.append(section)
        
        return sections

    def generate_html_documentation(self) -> str:
        """Generate HTML documentation."""
        # Generate sections
        tutorial_sections = self.generate_tutorial_sections() if self.config.include_tutorials else []
        api_sections = self.generate_api_reference_sections() if self.config.include_api_reference else []
        
        # Load template
        template_str = self._get_html_template()
        template = self.template_env.from_string(template_str)
        
        # Generate CSS based on theme
        css = self._generate_theme_css()
        
        return template.render(
            config=self.config,
            tutorial_sections=tutorial_sections,
            api_sections=api_sections,
            css=css,
            openapi_spec=self.openapi_spec,
            generation_time=datetime.now().isoformat()
        )

    def _get_html_template(self) -> str:
        """Get HTML template."""
        return '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ config.company_name }} {{ config.product_name }} Documentation</title>
    <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
    <link href="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-tomorrow.min.css" rel="stylesheet">
    <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/prism.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-python.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-javascript.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-bash.min.js"></script>
    <style>
        {{ css | safe }}
        {% if config.custom_css %}
        {{ config.custom_css | safe }}
        {% endif %}
    </style>
</head>
<body class="bg-gray-50 text-gray-900">
    <!-- Header -->
    <header class="bg-white shadow-sm border-b sticky top-0 z-50">
        <div class="max-w-7xl mx-auto px-4 sm:px-6 lg:px-8">
            <div class="flex justify-between items-center py-4">
                <div class="flex items-center space-x-3">
                    {% if config.logo_url %}
                    <img src="{{ config.logo_url }}" alt="{{ config.company_name }}" class="h-8 w-8">
                    {% endif %}
                    <div>
                        <h1 class="text-xl font-bold text-gray-900">{{ config.company_name }}</h1>
                        <p class="text-sm text-gray-600">{{ config.product_name }} v{{ config.version }}</p>
                    </div>
                </div>
                <nav class="hidden md:flex space-x-8">
                    <a href="#getting-started" class="text-gray-600 hover:text-gray-900">Getting Started</a>
                    <a href="#api-reference" class="text-gray-600 hover:text-gray-900">API Reference</a>
                    <a href="#tutorials" class="text-gray-600 hover:text-gray-900">Tutorials</a>
                </nav>
            </div>
        </div>
    </header>

    <!-- Main Content -->
    <main class="max-w-7xl mx-auto px-4 sm:px-6 lg:px-8 py-8">
        <div class="grid grid-cols-1 lg:grid-cols-4 gap-8">
            <!-- Sidebar -->
            <aside class="lg:col-span-1">
                <nav class="sticky top-24">
                    <h3 class="font-semibold text-lg mb-4">Table of Contents</h3>
                    <ul class="space-y-2">
                        {% if tutorial_sections %}
                        <li><a href="#getting-started" class="text-blue-600 hover:text-blue-800">Getting Started</a></li>
                        {% endif %}
                        {% if api_sections %}
                        <li><a href="#api-reference" class="text-blue-600 hover:text-blue-800">API Reference</a></li>
                        {% endif %}
                        {% if tutorial_sections %}
                        <li><a href="#tutorials" class="text-blue-600 hover:text-blue-800">Tutorials</a></li>
                        {% endif %}
                    </ul>
                </nav>
            </aside>

            <!-- Content -->
            <div class="lg:col-span-3">
                {% if tutorial_sections %}
                <section id="getting-started" class="mb-12">
                    <h2 class="text-3xl font-bold mb-6">Getting Started</h2>
                    {% for section in tutorial_sections %}
                    <div class="mb-8">
                        <h3 class="text-2xl font-semibold mb-4">{{ section.title }}</h3>
                        <div class="prose max-w-none">
                            {{ section.content | markdown | safe }}
                        </div>
                        {% if section.subsections %}
                        <div class="space-y-6 mt-6">
                            {% for subsection in section.subsections %}
                            <div class="border-l-4 border-blue-500 pl-4">
                                <h4 class="text-xl font-semibold mb-3">{{ subsection.title }}</h4>
                                <div class="prose max-w-none">
                                    {{ subsection.content | markdown | safe }}
                                </div>
                                {% if subsection.code_examples %}
                                <div class="mt-4 space-y-4">
                                    {% for example in subsection.code_examples %}
                                    <div class="bg-gray-900 rounded-lg p-4">
                                        <div class="flex justify-between items-center mb-2">
                                            <span class="text-sm font-mono text-gray-400">{{ example.language }}</span>
                                            <button class="copy-btn text-xs px-2 py-1 bg-gray-700 text-gray-300 rounded hover:bg-gray-600"
                                                    data-code="{{ example.code | replace('"', '&quot;') }}">
                                                Copy
                                            </button>
                                        </div>
                                        <pre><code class="language-{{ example.language }}">{{ example.code }}</code></pre>
                                    </div>
                                    {% endfor %}
                                </div>
                                {% endif %}
                            </div>
                            {% endfor %}
                        </div>
                        {% endif %}
                    </div>
                    {% endfor %}
                </section>
                {% endif %}

                {% if api_sections %}
                <section id="api-reference" class="mb-12">
                    <h2 class="text-3xl font-bold mb-6">API Reference</h2>
                    {% for section in api_sections %}
                    <div class="mb-12">
                        <h3 class="text-2xl font-semibold mb-6">{{ section.title }}</h3>
                        <div class="prose max-w-none">
                            {{ section.content | markdown | safe }}
                        </div>
                        {% if section.subsections %}
                        <div class="space-y-8 mt-6">
                            {% for subsection in section.subsections %}
                            <div class="bg-white rounded-lg shadow-md p-6 border">
                                <div class="flex items-center justify-between mb-4">
                                    <h4 class="text-xl font-semibold">
                                        <span class="method-badge method-{{ subsection.metadata.method | lower }}">{{ subsection.metadata.method }}</span>
                                        <code class="text-lg font-mono">{{ subsection.metadata.path }}</code>
                                    </h4>
                                </div>
                                <div class="prose max-w-none">
                                    {{ subsection.content | safe }}
                                </div>
                            </div>
                            {% endfor %}
                        </div>
                        {% endif %}
                    </div>
                    {% endfor %}
                </section>
                {% endif %}
            </div>
        </div>
    </main>

    <!-- Footer -->
    <footer class="bg-white border-t mt-12">
        <div class="max-w-7xl mx-auto px-4 sm:px-6 lg:px-8 py-8">
            <div class="text-center text-gray-600">
                <p>&copy; 2024 {{ config.company_name }}. All rights reserved.</p>
                <p class="mt-2">Generated on {{ generation_time | format_date }}</p>
            </div>
        </div>
    </footer>

    <script>
        // Copy functionality
        document.querySelectorAll('.copy-btn').forEach(btn => {
            btn.addEventListener('click', () => {
                navigator.clipboard.writeText(btn.dataset.code);
                btn.textContent = 'Copied!';
                setTimeout(() => {
                    btn.textContent = 'Copy';
                }, 2000);
            });
        });

        // Smooth scrolling
        document.querySelectorAll('a[href^="#"]').forEach(anchor => {
            anchor.addEventListener('click', function (e) {
                e.preventDefault();
                const target = document.querySelector(this.getAttribute('href'));
                if (target) {
                    target.scrollIntoView({
                        behavior: 'smooth',
                        block: 'start'
                    });
                }
            });
        });
    </script>
</body>
</html>
        '''

    def _generate_theme_css(self) -> str:
        """Generate CSS based on theme."""
        themes = {
            DocumentationTheme.DEFAULT: '''
                .method-badge { padding: 4px 8px; border-radius: 4px; font-weight: bold; font-size: 12px; }
                .method-get { background-color: #10b981; color: white; }
                .method-post { background-color: #3b82f6; color: white; }
                .method-put { background-color: #f59e0b; color: white; }
                .method-delete { background-color: #ef4444; color: white; }
                .params-table { width: 100%; border-collapse: collapse; }
                .params-table th, .params-table td { border: 1px solid #e5e7eb; padding: 8px; text-align: left; }
                .params-table th { background-color: #f9fafb; font-weight: 600; }
                .response-item { margin-bottom: 16px; padding: 16px; border: 1px solid #e5e7eb; border-radius: 8px; }
                .endpoint-description { margin-bottom: 24px; }
            ''',
            DocumentationTheme.DARK: '''
                body { background-color: #1f2937; color: #f9fafb; }
                .bg-white { background-color: #374151; }
                .text-gray-900 { color: #f9fafb; }
                .text-gray-600 { color: #d1d5db; }
                .border { border-color: #4b5563; }
                .method-badge { padding: 4px 8px; border-radius: 4px; font-weight: bold; font-size: 12px; }
                .method-get { background-color: #10b981; color: white; }
                .method-post { background-color: #3b82f6; color: white; }
                .method-put { background-color: #f59e0b; color: white; }
                .method-delete { background-color: #ef4444; color: white; }
            ''',
            DocumentationTheme.MODERN: '''
                .method-badge { padding: 6px 12px; border-radius: 20px; font-weight: 600; font-size: 14px; letter-spacing: 0.5px; }
                .method-get { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; }
                .method-post { background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); color: white; }
                .method-put { background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); color: white; }
                .method-delete { background: linear-gradient(135deg, #fa709a 0%, #fee140 100%); color: white; }
                .params-table { width: 100%; border-collapse: collapse; border-radius: 8px; overflow: hidden; }
                .params-table th, .params-table td { border: none; padding: 12px; }
                .params-table th { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; }
                .response-item { margin-bottom: 20px; padding: 20px; border: none; border-radius: 12px; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1); }
            '''
        }
        
        return themes.get(self.config.theme, themes[DocumentationTheme.DEFAULT])

    def generate_markdown_documentation(self) -> str:
        """Generate Markdown documentation."""
        sections = []
        
        # Add title and introduction
        sections.append(f"# {self.config.company_name} {self.config.product_name} Documentation\n")
        sections.append(f"Version {self.config.version}\n")
        sections.append("Generated on " + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "\n")
        
        # Add getting started
        if self.config.include_tutorials:
            tutorial_sections = self.generate_tutorial_sections()
            sections.append("## Getting Started\n")
            for section in tutorial_sections:
                sections.append(f"### {section.title}\n")
                sections.append(f"{section.content}\n")
                
                for subsection in section.subsections:
                    sections.append(f"#### {subsection.title}\n")
                    sections.append(f"{subsection.content}\n")
                    
                    for example in subsection.code_examples:
                        sections.append(f"```{example['language']}\n{example['code']}\n```\n")
        
        # Add API reference
        if self.config.include_api_reference:
            api_sections = self.generate_api_reference_sections()
            sections.append("## API Reference\n")
            
            for section in api_sections:
                sections.append(f"### {section.title}\n")
                sections.append(f"{section.content}\n")
                
                for subsection in section.subsections:
                    sections.append(f"#### {subsection.metadata['method']} {subsection.metadata['path']}\n")
                    sections.append(f"{subsection.content}\n")
        
        return "\n".join(sections)

    def generate_pdf_documentation(self, html_content: str) -> bytes:
        """Generate PDF documentation."""
        options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'no-outline': None
        }
        
        return pdfkit.from_string(html_content, False, options=options)

    def generate_word_documentation(self, markdown_content: str) -> bytes:
        """Generate Word documentation."""
        doc = Document()
        
        # Add title
        doc.add_heading(f'{self.config.company_name} {self.config.product_name} Documentation', 0)
        
        # Add metadata
        p = doc.add_paragraph()
        p.add_run(f"Version {self.config.version}")
        p.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Parse and add markdown content
        lines = markdown_content.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('```'):
                # Handle code blocks (simplified)
                continue
            elif line.strip():
                doc.add_paragraph(line)
        
        return doc.save()

    def generate_documentation(self) -> Dict[str, str]:
        """Generate documentation in all specified formats."""
        logger.info(f"Generating documentation in formats: {[f.value for f in self.config.formats]}")
        
        generated_files = {}
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Generate HTML
        if OutputFormat.HTML in self.config.formats:
            html_content = self.generate_html_documentation()
            html_file = Path(self.config.output_dir) / f"documentation_{timestamp}.html"
            with open(html_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            generated_files['html'] = str(html_file)
            logger.info(f"HTML documentation generated: {html_file}")
        
        # Generate Markdown
        if OutputFormat.MARKDOWN in self.config.formats:
            markdown_content = self.generate_markdown_documentation()
            md_file = Path(self.config.output_dir) / f"documentation_{timestamp}.md"
            with open(md_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            generated_files['markdown'] = str(md_file)
            logger.info(f"Markdown documentation generated: {md_file}")
        
        # Generate PDF (requires HTML)
        if OutputFormat.PDF in self.config.formats and OutputFormat.HTML in self.config.formats:
            pdf_content = self.generate_pdf_documentation(html_content)
            pdf_file = Path(self.config.output_dir) / f"documentation_{timestamp}.pdf"
            with open(pdf_file, 'wb') as f:
                f.write(pdf_content)
            generated_files['pdf'] = str(pdf_file)
            logger.info(f"PDF documentation generated: {pdf_file}")
        
        # Generate Word (requires Markdown)
        if OutputFormat.WORD in self.config.formats and OutputFormat.MARKDOWN in self.config.formats:
            word_content = self.generate_word_documentation(markdown_content)
            word_file = Path(self.config.output_dir) / f"documentation_{timestamp}.docx"
            generated_files['word'] = str(word_file)
            logger.info(f"Word documentation generated: {word_file}")
        
        # Generate JSON
        if OutputFormat.JSON in self.config.formats:
            json_data = {
                'metadata': {
                    'company': self.config.company_name,
                    'product': self.config.product_name,
                    'version': self.config.version,
                    'generated_at': datetime.now().isoformat(),
                    'formats': [f.value for f in self.config.formats]
                },
                'openapi_spec': self.openapi_spec,
                'tutorial_sections': self.generate_tutorial_sections() if self.config.include_tutorials else [],
                'api_sections': self.generate_api_reference_sections() if self.config.include_api_reference else []
            }
            json_file = Path(self.config.output_dir) / f"documentation_{timestamp}.json"
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, indent=2, ensure_ascii=False)
            generated_files['json'] = str(json_file)
            logger.info(f"JSON documentation generated: {json_file}")
        
        return generated_files


# CLI usage
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Generate documentation")
    parser.add_argument("--openapi-spec", default="docs/openapi_comprehensive.yaml", help="OpenAPI spec path")
    parser.add_argument("--output-dir", default="generated_docs", help="Output directory")
    parser.add_argument("--formats", nargs="+", default=["html"], help="Output formats")
    parser.add_argument("--theme", default="default", help="Documentation theme")
    parser.add_argument("--no-tutorials", action="store_true", help="Exclude tutorials")
    parser.add_argument("--no-api-ref", action="store_true", help="Exclude API reference")
    parser.add_argument("--custom-css", help="Custom CSS file")
    parser.add_argument("--company", default="RaptorFlow", help="Company name")
    parser.add_argument("--product", default="Backend API", help="Product name")
    parser.add_argument("--version", default="1.0.0", help="Documentation version")
    
    args = parser.parse_args()
    
    # Create configuration
    config = DocumentationConfig(
        openapi_spec_path=args.openapi_spec,
        output_dir=args.output_dir,
        formats=[OutputFormat(f) for f in args.formats],
        theme=DocumentationTheme(args.theme),
        include_tutorials=not args.no_tutorials,
        include_api_reference=not args.no_api_ref,
        custom_css=open(args.custom_css).read() if args.custom_css else None,
        company_name=args.company,
        product_name=args.product,
        version=args.version
    )
    
    # Generate documentation
    generator = DocumentationGenerator(config)
    generated_files = generator.generate_documentation()
    
    print(f"\nDocumentation generated successfully:")
    for format_type, file_path in generated_files.items():
        print(f"  {format_type.upper()}: {file_path}")
